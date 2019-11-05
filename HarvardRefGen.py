"""
Manchester Harvard Reference Generator

This code will create GUI windows to select a source to format in Manchester Harvard
Reference style. Fields will be created to enter the appropriate information. When genereate
reference is chosen the code will append a file called references.docx with numbered and
formatted references generated from the inputs. Ensure the file is not open while attempting
to generate references. Also ensure you follow the example entry formatting, only putting
punctuation if shown in the example input and capitalising as shown in the example input
to make the output correctly formatted.

ISBN search is a trial feature that allows you to search the ISBN of a book and the software
will attempt to autofill the formatted values. This will not always be 100% accurate so it
is recommended that you check the fetched data to ensure it is accurate and correctly
formatted for use in generating the reference. Data is fetched from google books and
formatted before entry but if the google books metadata is incorrectly stored online then
formatting will not be correct. This feature requires internet connection to function.

A lot of this code is very repetitive in functions so one example of each will be fully
commented to explain proccesses but in interest of comment efficiency not all will be.

Joshua Utting

30/08/2019

"""

#imports list
from tkinter import *
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
import requests
import numpy as np
from PIL import Image,ImageTk
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from matplotlib.figure import Figure

try: #attempt to open a file called references.docx
    document = Document('references.docx')
    length = len(document.paragraphs) #if file opens find the number of paragraphs (i.e. references)   
    count = length #set count to start from one after the last number in the list
except:
    document = Document() #if file does not exist create a new file
    count=0 #set count to 0 to start reference numbering from 1

#define word document text style as Times New Roman 12pt
style = document.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

def hasNumbers(inputString):
    return any(char.isdigit() for char in inputString) #check if there is a number in the inputstring

def suffix(d):
    return 'th' if 11<=d<=13 else {1:'st',2:'nd',3:'rd'}.get(d%10, 'th') #for the edition number find the appropriate suffix

#function to search an entered isbn number and autofill the form basedon webscraped data
def isbnsearch(isbnvar,authorvar,yearvar,titlevar,subtitlevar,editionvar,pubplacevar,publishervar):
    try: #provide a catch if this fails anywhere
        a = np.array([]) #create array of links
        isbn = isbnvar.get() #get the entered isbn string
        urlstring = ('https://www.google.com/search?tbm=bks&q=' + isbn) #use url of google book search of isbn value
        #send request to the site also sending a header to be treated as a headed browser
        response = requests.get(urlstring, headers={'user-agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_13_2) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/63.0.3239.132 Safari/537.36'})
        html1 = response.text #store the source code of the page as a string
        soup = BeautifulSoup(html1, 'html.parser') #parse the string using a html parser
        
        results = soup.find_all("div",attrs={"class":"r"}) #find all div tags with the class r (this tag stores results of search)
        for result in results:
            links = result.find_all('a') #find any instance of a link tag in each result of first search
            for link in links:
                a = np.append(a,link.get('href')) #for each tag store the link string
    
        for i in range(0,a.size-1): #search the array of links
            if(isbn in str(a[i])): 
                linkno = i #if the searched isbn is in the link then store which link to use, used to skip ads
                break #break loop when found
            else:
                pass

        bookurl = a[linkno] #store new url string as link to the book's page on google books
        #send request to the site also sending a header to be treated as a headed browser
        bookresponse = requests.get(bookurl, headers={'user-agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_13_2) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/63.0.3239.132 Safari/537.36'})
        html2 = bookresponse.text #store the source code of this page as a string
        soup = BeautifulSoup(html2, 'html.parser') #parse this string using a html parser
        
        try: #check to see if in the about page of the book or on the first page, need about page for metadata
            results1 = soup.find('a',attrs={'id':'sidebar-atb-link'}) #search for any links stored in a sidebar, should not find any if on about page
            newlink = results1.get('href') #if instance found of sidebar link get the link of this about page
            #send request to the site also sending a header to be treated as a headed browser
            bookresponse2 = requests.get(newlink, headers={'user-agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_13_2) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/63.0.3239.132 Safari/537.36'})
            html3 = bookresponse2.text #store the source code of this page as a string
            soup = BeautifulSoup(html3, 'html.parser') #parse the string using a html parser
        except:
            soup = BeautifulSoup(html2, 'html.parser') #if on the about page already parse the original page using html parser
    
        results = soup.find_all('td',attrs={'class':'metadata_label'}) #search for all td tags with class metadata_label
    
        metadatalabels = np.array([]) #create empty array of metadatalabels
        metadatavalues = np.array([]) #create empty array of metadatavalues
        for result in results:
            metadatalabels = np.append(metadatalabels,result.text) #store the text of all metadatalabels in array
        metadatalabels = metadatalabels[:-1] #remove final label (export citation label)
        results = soup.find_all('td',attrs={'class':'metadata_value'}) #search for all td tags with class metadata_value
        for result in results:
            datavals = result.find_all('span',attrs={'dir':'ltr'}) #search for all span tags with dir=ltr in these td classes
            for val in datavals:
                metadatavalues = np.append(metadatavalues,(val.text)) #store the text of the metadatavalues in array
        metadatavalues = metadatavalues[:-3] #remove the last three metadatavalues (three options to export citation)            
        
        for i in range(0,len(metadatavalues)-1): #loop through metadatavalues
            if(hasNumbers(metadatavalues[i])==True): #if numbers in element
                attempt = metadatavalues[i].split(',') #try splitting the element by comma
                attempt[1] = attempt[1].lstrip() #remove the white space at the start of second element
                if(len(attempt[1])==4): #if the second element of the attempt is a year
                    success = i #store the value
                    break #break loop
                else:
                    pass
            else:
                pass
        
        if('Edition' not in metadatalabels): #if edition is missing
            metadatalabels = metadatalabels[:-3] #remove the next three metadatalabels
            metadatavalues = metadatavalues[0:success+1] #replace metadatavalues with values from 0 to the success count
        else: #if edition is present
            metadatalabels = metadatalabels[:-2] #remove the next two metadatalabels
            metadatavalues = metadatavalues[0:success+1] #replace metadatavalues with values from 0 to the success count
        
        if(len(metadatalabels)==3): #check if 3 labels, meaning eidtion not found
            numauthors = int(len(metadatavalues)-len(metadatalabels)+1) #number of authors is the difference between the number of values and number of labels + 1
            #extra note on above: author is the only one to have multiple values for a single label hence subtitution method works
            title = str(metadatavalues[0]).lower() #set title variable as the first metadatavalue in lowercase
            title = title.capitalize() #capitalize first letter of title
            if(':' in title): #check to see if there is a semicolon in the title (indicating subtitle)
                title2 = title.split(':') #split the title at the semicolon
                title = title2[0] #overwrite title as the first element of the split
                subtitle = str(title2[1]).lower() #store subtitle as second element of split in lowercase
                subtitle = subtitle.lstrip() #remove the space at the start of the subtitle string
            else:
                subtitle='' #if no colon set subtitle as blank
            authors = metadatavalues[1:numauthors+1] #set authors as the second metadatavalue to the element of number of authors + 1 (as the array started from 0)
            edn = '' #set edition as blank as it is not listed
            pubandyear = metadatavalues[numauthors+1] #publisher and year listed in the same value, set this as the element after authors as edition tag is missing
            
            try:
                pubyearsplit = pubandyear.split(',') #attempt to split the publisher and published year by the comma
                publisher = pubyearsplit[0] #set the publisher as the first element of the split
                year = pubyearsplit[1].lstrip() #set the published year as the second element of the strip removing space at start
            except:
                publisher = pubandyear #if fails set publisher to the original metadatavalue
                year='' #set year to blank
        else: #if there is not 6 metadatalabels (empirically means there will be 7)
            for i in range(0,len(metadatavalues)): #complicated loop process over the number of metadatavalues
                if(metadatavalues[i]=='illustrated' or hasNumbers(metadatavalues[i])==True): #empirically the edition will always be listed as illustrated or be the first tag to have a number in it and will be the tag after authors
                    #continuation of above: as a results this value can be used to determine the number of authors
                    numauthors = int(len(metadatavalues[1:i])) #number of authors is the number of metadatavalues after title and beofre this edition value
                    title = str(metadatavalues[0]).lower() #set title as the first metadatavalue in lowercase
                    title = title.capitalize() #capitalize first letter of title string
                    if(':' in title): #check to see if there is a semicolon in the title (indicating subtitle)
                        title2 = title.split(':') #split the title at the semicolon
                        title = title2[0] #overwrite title as the first element of the split
                        subtitle = str(title2[1]).lower() #store subtitle as second element of split in lowercase
                        subtitle = subtitle.lstrip() #remove the space at the start of the subtitle string
                    else:
                        subtitle='' #if no colon set subtitle as blank
                    authors = metadatavalues[1:numauthors+1] #set authors as the second metadatavalue to the element of number of authors + 1 (as the array started from 0)
                    print(authors)
                    edn = metadatavalues[numauthors+1] #set edition as the metadatavalue following the final author
                    try: #check a case for if the edition has a number and unrelated text string like 'illustrated'
                        splitup = edn.split(',') #these options will be comma separated so split at a comma
                        for i in range(0,len(splitup)): #loop over number of items in split
                            if(hasNumbers(splitup[i])==True): #check if the element has a number in it
                                edn = splitup[i] #if yes store edition as the text of that element from the splitup (now a number)
                                break #break loop to carry on
                            else:
                                pass
                        try:
                            ending = suffix(int(edn)) #find a suffix for the edition based on the suffix function (e.g. th,nd,rd)
                            edn = edn+ending+' edn' #store edition as the number plus suffic plus ' edn'
                        except:
                            edn = '1st edn' #if edition still came back as illustrated replace it with 1st edition
                        
                    except:
                        pass #if this check fails leave as the scraped value
                    pubandyear = metadatavalues[numauthors+2] #store variable publisher and published year as metadatavalues from the final author + 2
                    
                    try:
                        pubyearsplit = pubandyear.split(',') #attempt to split publisher and published year string at comma as they are comma separated
                        publisher = pubyearsplit[0] #set publisher as first element of the split
                        year = pubyearsplit[1].lstrip() #set published year as the second element of the split and remove space at start of string
                    except:
                        publisher = pubandyear #if fails store publisher as the original scraped string
                        year='' #set year to blank
                    else:
                        pass
                    break #break the loop as all useful values are stored somewhere now
                else:
                    pass #if not on edition value then pass
        
        newauthor = [] #set blank list for formatted authors
        
        if(numauthors==1): #if there is 1 author
            try:
                authorslist = authors #store a copy incase of fail
                author1 = str(authorslist[0]).split(' ') #split the author string by space to get name split in to elements
                l1 = len(author1) #store number of elements in author name
                newauthor.append(author1[l1-1]+ ', ' + author1[l1-2][0]+'.') #append newauthor list as '[author surname], [intiial].'
                formattedauthor = newauthor[0] #set formattedauthor as the first element of newauthor
            except:
                formattedauthor = authors #if fails store formnattedauthor as scraped author value
        elif(numauthors==2): #if there is 2 authors
            try:
                authorslist = authors #store a copy incase of fail
                author1 = str(authorslist[0]).split(' ') #split the first author's name in list by space
                l1 = len(author1) #store number of elements in first author name
                author2 = str(authorslist[1]).split(' ') #split the second author's name in list by space
                l2 = len(author2) #store number of elements in second author name
                newauthor.append(author1[l1-1]+ ', ' + author1[l1-2][0]+'.') #append newauthor list as '[author 1 surname], [intiial].'
                newauthor.append(author2[l2-1]+ ', ' + author2[l2-2][0]+'.') #append newauthor list as '[author 2 surname], [intiial].'
                formattedauthor = newauthor[0]+' and '+newauthor[1] #store formatted author as first element of newauthor ' and ' second element of newauthor
            except:
                formattedauthor=authors #if fails store formattedauthor as scraped author value
        elif(numauthors>=3): #if 3 or more authors
            try:
                formattedauthor='' #try setting formattedauthor blank
                authorslist = authors #store a copy incase of fail
                alist = [] #create authorlist
                lt = [] #create author name length list
                for i in range(0,numauthors): #loop over the number of authors
                    alist.append(str(authorslist[i]).split(' ')) #append author list with space split author names
                    lt.append(len(str(authorslist[i]).split(' '))) #append author name length list with number of elements in author name
                    newauthor.append(alist[i][lt[i]-1]+', '+alist[i][lt[i]-2][0]+'.') #append newauthor with each '[author surname], [intial].'
                for i in range(0,len(newauthor)-1): #loop over all but the final author
                    formattedauthor = formattedauthor+str(newauthor[i])+', ' #add each new author to formattedauthor string with a comma and space after
                formattedauthor = str(formattedauthor).rstrip(', ')+' and '+str(newauthor[len(newauthor)-1]) #to finish the string remove the last ', ' and replace with ' and ' [final newauthor element]
            except:
                formattedauthor = authors #if fails store formattedauthor as scraped authos value
        
        try:
            pubsearch = publisher #last unfound value is publisher location not found on about page, store a copy incase of fail
            pubsearch = pubsearch.replace(' ','_') #replace any spaces in publisher string with _
            url='https://en.wikipedia.org/wiki/'+pubsearch #store url string as wikipedia page of publisher
            #send request to the site also sending a header to be treated as a headed browser
            response = requests.get(url, headers={'user-agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_13_2) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/63.0.3239.132 Safari/537.36'})
            html4 = response.text #store page source as string
            soup = BeautifulSoup(html4, 'html.parser') #parse string with html parser
            results2 = soup.find("td",attrs={"class":"label"}) #find all instances of td tag with the class label
            place = results2.text #get the text of the contents of this tag for the publisher location
            newplace = place.split(',') #split publisher location on comma as it will usually give format city,state,country
            pubplace = newplace[0] #store publisher location as city element
        except:
            pubplace = '' #if fails store publisher location as blank
        
        try:
            titlevar.set(title) #set the title variable of entry form to the title string
        except:
            pass
        try:
            yearvar.set(year) #set the year variable of entry form to the year string
        except:
            pass
        try:
            authorvar.set(formattedauthor) #set the author variable of the entry form to the formattedauthor string
        except:
            pass
        try:
            editionvar.set(edn) #set the edition variable of the entry field to the edn string
        except:
            pass
        try:
            publishervar.set(publisher) #set the publisher variable of the entry field to the publisher string
        except:
            pass
        try:
            subtitlevar.set(subtitle) #set the subtitle variable of the entry field to the subtitle string
        except:
            pass
        try:
            pubplacevar.set(pubplace) #set the publisher location variable of the entry field to the pubplace string
        except:
            pass
    except: #if isbn search has failed
        isbnfail = Tk() #open a tkinter window
        #create text label informing search has failed and to try again, pack in window
        infolabel = Label(isbnfail,text='ISBN not found. If this is your first time seeing this message please try again.\nTry with and without a hyphen after the first three digits.\nIf the problem persists on numerous attempts then Google Books does not correctly store the metadata so the search will not be able to autofill.').pack()
        okbutton = Button(isbnfail,text='Ok',command=lambda: isbnfail.destroy()).pack() #create ok button to close window and pack in window
        isbnfail.mainloop() #keep window displaying


        
    
#define function to get the choice from the drop down and call the appropriate reference window
def getchoice(*choices):
    global reftpyechoice, reftypevar, count
    reftypechoice = reftypevar.get() #get the value of the drop down
    if(reftypechoice=='Book'): #if the value is book etc.
        bookwin() #call the bookwindow function to show the entry fields for a book etc
    elif(reftypechoice=='Website'):
        webwin()
    elif(reftypechoice=='Journal (printed)'):
        journal1win()
    elif(reftypechoice=='Journal (electronic)'):
        journal2win()
    elif(reftypechoice=='Lecturer Handout'):
        lechandoutwin()
    elif(reftypechoice=='Thesis'):
        thesiswin()
    elif(reftypechoice=='Presentation'):
        presentwin()
    elif(reftypechoice=='Chapter From Edited Book'):
        editbookchapterwin()
    elif(reftypechoice=='Edited Book'): 
        editedbookwin()
    elif(reftypechoice=='Software'): 
        softwarewin()
    elif(reftypechoice=='Illustration'): 
        illustrationwin()
    elif(reftypechoice=='E-Book'):
        ebookwin()
    elif(reftypechoice=='Report From Organisation'):
        organisationreportwin()
    elif(reftypechoice=='Self-Citation'):
        selfcitationwin()
    elif(reftypechoice=='Government/Corporate Publications'):
        govpubswin()
    elif(reftypechoice=='Custom'):
        customwin()
    
#define the book genating reference function
#these generation functions are called by the relevant sourcewin functions below, e.g. bookwin()
def bookgen(authorstring,yearstring,titlestring,subtitlestring,editionstring,pubplacestring,publisherstring,referenceframe):
    global count
    count=count+1 #add one to the count to increase the number in the reference list
    #get the values of each of the inputs of the book window entry fields
    author = authorstring.get()
    year = yearstring.get()
    title = titlestring.get()
    subtitle = subtitlestring.get()
    edition = editionstring.get()
    pubplace = pubplacestring.get()
    publisher = publisherstring.get()

    if(subtitle!=''): #check if there has been an entry in the subtitle window, if so:
        newsubtitlestring = subtitle + '.' #add full stop to the end of subtitle
        subtitle = newsubtitlestring #overwrite subtitle
        newtitlestring = title + ': ' #add colon to the end of title
        title = newtitlestring+subtitle #new title is in format title: subtitle.
    else:
        newtitlestring = title + '.' #if subtitle is empty add full stop to the end of title 
        title = newtitlestring #new title is in format title.
    
    p = document.add_paragraph("[%s] " % count) #add a new paragraph to document starting with count value e.g reference [3]
    p.style = document.styles['Normal'] #write paragraph in previously defined font style
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY #justify aligning of paragraph as per required for UoM reports
    p.add_run("%s (%s). " % (author,year)) #add the title and year in brackets to the reference
    p.add_run(title).italic = True #add the title to the paragraph in italics
    p.add_run(" %s. %s: %s." % (edition,pubplace,publisher)) #add the edition. publishier location: publisher name. to the paragraph
    document.save('references.docx') #save the appending to references.docx
    success = Label(referenceframe,text=('Reference [%s] Successfully Generated' % count)).grid(row=9,column=0,columnspan=2) #create a label at the bottom of the window informing reference [number] added successfully

#define the web reference generating function
#works as bookgen but all formatting matches that required for websites
def webgen(authorstring,yearstring,titlestring,urlstring,accessstring,referenceframe):
    global count
    count=count+1
    author = authorstring.get()
    year = yearstring.get()
    title = titlestring.get()
    url = urlstring.get()
    access = accessstring.get()
    
    p = document.add_paragraph("[%s] " % count)
    p.style = document.styles['Normal']
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run("%s (%s). " % (author,year))
    p.add_run(title).italic = True
    p.add_run(". Available at: %s, (Accessed: %s)." % (url,access))
    document.save('references.docx')
    success = Label(referenceframe,text=('Reference [%s] Successfully Generated' % count)).grid(row=8,column=0,columnspan=2)
    
#define the printed journal reference generating function
#works as bookgen but all formatting matches that required for printed journals
def journal1gen(authorstring,yearstring,titlearticlestring,titlestring,volumestring,issuestring,pagestring,referenceframe):
    global count
    count=count+1
    author = authorstring.get()
    year = yearstring.get()
    titlearticle = titlearticlestring.get()
    title = titlestring.get()
    volume = volumestring.get()
    issue = issuestring.get()
    page = pagestring.get()
    
    p = document.add_paragraph("[%s] " % count)
    p.style = document.styles['Normal']
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run("%s (%s). ‘%s’, " % (author,year,titlearticle))
    p.add_run(title).italic = True
    p.add_run(", %s(%s), pp. %s." % (volume,issue,page))
    document.save('references.docx')
    success = Label(referenceframe,text=('Reference [%s] Successfully Generated' % count)).grid(row=8,column=0,columnspan=2)

#define the electronic journal reference generating function
#works as bookgen but all formatting matches that required for electronic journals
def journal2gen(authorstring,yearstring,titlearticlestring,titlestring,volumestring,issuestring,pagestring,urlstring,accessstring,referenceframe):
    global count
    count=count+1
    author = authorstring.get()
    year = yearstring.get()
    titlearticle = titlearticlestring.get()
    title = titlestring.get()
    volume = volumestring.get()
    issue = issuestring.get()
    page = pagestring.get()
    url = urlstring.get()
    access = accessstring.get()
    
    p = document.add_paragraph("[%s] " % count)
    p.style = document.styles['Normal']
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run("%s (%s). ‘%s’, " % (author,year,titlearticle))
    p.add_run(title).italic = True
    p.add_run(", %s(%s), pp. %s. [Online]. Available at: %s (Accessed: %s)." % (volume,issue,page,url,access))
    document.save('references.docx')
    success = Label(referenceframe,text=('Reference [%s] Successfully Generated' % count)).grid(row=10,column=0,columnspan=2)

#define the lecturer handout reference generating function
#works as bookgen but all formatting matches that required for lecturer handouts
def lechandoutgen(authorstring,yearstring,titlestring,codestring,modtitlestring,institutionstring,referenceframe):
    global count
    count=count+1
    author = authorstring.get()
    year = yearstring.get()
    title = titlestring.get()
    code = codestring.get()
    modtitle = modtitlestring.get()
    institution = institutionstring.get()
    
    unitname = code + ': ' + modtitle
    
    p = document.add_paragraph("[%s] " % count)
    p.style = document.styles['Normal']
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run("%s (%s). ‘%s’, " % (author,year,title))
    p.add_run(unitname).italic = True
    p.add_run(". %s. Unpublished." % (institution))
    document.save('references.docx')
    success = Label(referenceframe,text=('Reference [%s] Successfully Generated' % count)).grid(row=7,column=0,columnspan=2)

#define the thesis reference generating function
#works as bookgen but all formatting matches that required for theses
def thesisgen(authorstring,yearstring,titlestring,levelstring,institutionstring,referenceframe):
    global count
    count=count+1
    author = authorstring.get()
    year = yearstring.get()
    title = titlestring.get()
    level = levelstring.get()
    institution = institutionstring.get()
    
    p = document.add_paragraph("[%s] " % count)
    p.style = document.styles['Normal']
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run("%s (%s). " % (author,year))
    p.add_run(title).italic = True
    p.add_run(". %s. %s." % (level,institution))
    document.save('references.docx')
    success = Label(referenceframe,text=('Reference [%s] Successfully Generated' % count)).grid(row=6,column=0,columnspan=2)

#define the presentation reference generating function
#works as bookgen but all formatting matches that required for presentations
def presentgen(authorstring,yearstring,titlestring,codestring,modtitlestring,urlstring,accessstring,referenceframe):
    global count
    count=count+1
    author = authorstring.get()
    year = yearstring.get()
    title = titlestring.get()
    code = codestring.get()
    modtitle = modtitlestring.get()
    url = urlstring.get()
    access = accessstring.get()
    
    p = document.add_paragraph("[%s] " % count)
    p.style = document.styles['Normal']
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run("%s (%s). " % (author,year))
    p.add_run(title).italic = True
    p.add_run(" [presentation]. ")
    p.add_run("%s: %s" % (code,modtitle)).italic = True
    p.add_run(". Available at: %s (Accessed: %s)." % (url,access))
    document.save('references.docx')
    success = Label(referenceframe,text=('Reference [%s] Successfully Generated' % count)).grid(row=8,column=0,columnspan=2)    

#define the edited book reference generating function
#works as bookgen but all formatting matches that required for edited books
def editbookchaptergen(authorstring,yearstring,chapterstring,editorstring,titlestring,subtitlestring,pubplacestring,publisherstring,pagesstring,referenceframe):
    global count
    count=count+1
    author = authorstring.get()
    year = yearstring.get()
    chapter = chapterstring.get()
    editor = editorstring.get()
    title = titlestring.get()
    subtitle = subtitlestring.get()
    pubplace = pubplacestring.get()
    publisher = publisherstring.get()
    pages = pagesstring.get()
    
    noeds = int((1+len(editor.split(',')))/2)
    if(noeds>=2):
        edoreds = 'eds'
    else:
        edoreds = 'ed'

    if(subtitle!=''): #check if there has been an entry in the subtitle window, if so:
        newsubtitlestring = subtitle + '.' #add full stop to the end of subtitle
        subtitle = newsubtitlestring #overwrite subtitle
        newtitlestring = title + ': ' #add colon to the end of title
        title = newtitlestring+subtitle #new title is in format title: subtitle.
    else:
        newtitlestring = title + '.' #if subtitle is empty add full stop to the end of title 
        title = newtitlestring #new title is in format title.
    
    p = document.add_paragraph("[%s] " % count) #add a new paragraph to document starting with count value e.g reference [3]
    p.style = document.styles['Normal'] #write paragraph in previously defined font style
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY #justify aligning of paragraph as per required for UoM reports
    p.add_run("%s (%s). '%s', in %s (%s.) " % (author,year,chapter,editor,edoreds)) #add the title and year in brackets to the reference
    p.add_run(title).italic = True #add the title to the paragraph in italics
    p.add_run(" %s: %s, pp %s." % (pubplace,publisher,pages)) #add the edition. publishier location: publisher name. to the paragraph
    document.save('references.docx') #save the appending to references.docx
    success = Label(referenceframe,text=('Reference [%s] Successfully Generated' % count)).grid(row=11,column=0,columnspan=2) #create a label at the bottom of the window informing reference [number] added successfully

def editedbookgen(authorstring,yearstring,titlestring,subtitlestring,editionstring,pubplacestring,publisherstring,referenceframe):
    global count
    count=count+1 #add one to the count to increase the number in the reference list
    #get the values of each of the inputs of the book window entry fields
    author = authorstring.get()
    year = yearstring.get()
    title = titlestring.get()
    subtitle = subtitlestring.get()
    edition = editionstring.get()
    pubplace = pubplacestring.get()
    publisher = publisherstring.get()

    noeds = int((1+len(author.split(',')))/2)
    if(noeds>=2):
        edoreds = 'eds'
    else:
        edoreds = 'ed'

    if(subtitle!=''): #check if there has been an entry in the subtitle window, if so:
        newsubtitlestring = subtitle + '.' #add full stop to the end of subtitle
        subtitle = newsubtitlestring #overwrite subtitle
        newtitlestring = title + ': ' #add colon to the end of title
        title = newtitlestring+subtitle #new title is in format title: subtitle.
    else:
        newtitlestring = title + '.' #if subtitle is empty add full stop to the end of title 
        title = newtitlestring #new title is in format title.
    
    p = document.add_paragraph("[%s] " % count) #add a new paragraph to document starting with count value e.g reference [3]
    p.style = document.styles['Normal'] #write paragraph in previously defined font style
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY #justify aligning of paragraph as per required for UoM reports
    p.add_run("%s (%s.) (%s). " % (author,edoreds,year)) #add the title and year in brackets to the reference
    p.add_run(title).italic = True #add the title to the paragraph in italics
    if(edition!=''):
        p.add_run(" %s. %s: %s." % (edition,pubplace,publisher)) #add the edition. publishier location: publisher name. to the paragraph    
    else:
        p.add_run(" %s: %s." % (pubplace,publisher)) #add the edition. publishier location: publisher name. to the paragraph
    document.save('references.docx') #save the appending to references.docx
    success = Label(referenceframe,text=('Reference [%s] Successfully Generated' % count)).grid(row=9,column=0,columnspan=2) #create a label at the bottom of the window informing reference [number] added successfully

def softwaregen(authorstring,yearstring,titlestring,versionstring,urlstring,accessstring,referenceframe):
    global count
    count=count+1
    author = authorstring.get()
    year = yearstring.get()
    title = titlestring.get()
    version = versionstring.get()
    url = urlstring.get()
    access = accessstring.get()
    
    p = document.add_paragraph("[%s] " % count)
    p.style = document.styles['Normal']
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run("%s (%s). " % (author,year))
    p.add_run(title).italic = True
    p.add_run(" (%s) [Computer Program]. Available at: %s (Accessed: %s)." % (version,url,access))
    document.save('references.docx')
    success = Label(referenceframe,text=('Reference [%s] Successfully Generated' % count)).grid(row=7,column=0,columnspan=2)

def illustrationgen(authorstring,yearstring,titlestring,subtitlestring,pubplacestring,publisherstring,pagestring,urlstring,accessstring,typestring,referenceframe):
    global count
    count=count+1 #add one to the count to increase the number in the reference list
    #get the values of each of the inputs of the book window entry fields
    author = authorstring.get()
    year = yearstring.get()
    title = titlestring.get()
    subtitle = subtitlestring.get()
    pubplace = pubplacestring.get()
    publisher = publisherstring.get()
    page = pagestring.get()
    url = urlstring.get()
    access = accessstring.get()
    typem = typestring.get()

    if(subtitle!=''): #check if there has been an entry in the subtitle window, if so:
        newsubtitlestring = subtitle + '.' #add full stop to the end of subtitle
        subtitle = newsubtitlestring #overwrite subtitle
        newtitlestring = title + ': ' #add colon to the end of title
        title = newtitlestring+subtitle #new title is in format title: subtitle.
    else:
        newtitlestring = title + '.' #if subtitle is empty add full stop to the end of title 
        title = newtitlestring #new title is in format title.

    p = document.add_paragraph("[%s] " % count) #add a new paragraph to document starting with count value e.g reference [3]
    p.style = document.styles['Normal'] #write paragraph in previously defined font style
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY #justify aligning of paragraph as per required for UoM reports
    p.add_run("%s (%s). " % (author,year)) #add the title and year in brackets to the reference
    p.add_run(title).italic = True #add the title to the paragraph in italics
    if(page!=''):
        p.add_run(" [%s]. %s: %s. pp. %s" % (typem,pubplace,publisher,page)) #add the edition. publishier location: publisher name. to the paragraph    
    else:
        p.add_run(" [%s]. Available at: %s (Accessed: %s)" % (typem,url,access)) #add the edition. publishier location: publisher name. to the paragraph
    document.save('references.docx') #save the appending to references.docx
    success = Label(referenceframe,text=('Reference [%s] Successfully Generated' % count)).grid(row=11,column=0,columnspan=2) #create a label at the bottom of the window informing reference [number] added successfully

def ebookgen(authorstring,yearstring,titlestring,subtitlestring,editionstring,providerstring,urlstring,referenceframe):
    global count
    count=count+1 #add one to the count to increase the number in the reference list
    #get the values of each of the inputs of the book window entry fields
    author = authorstring.get()
    year = yearstring.get()
    title = titlestring.get()
    subtitle = subtitlestring.get()
    edition = editionstring.get()
    provider = providerstring.get()
    url = urlstring.get()

    if(subtitle!=''): #check if there has been an entry in the subtitle window, if so:
        newsubtitlestring = subtitle + '.' #add full stop to the end of subtitle
        subtitle = newsubtitlestring #overwrite subtitle
        newtitlestring = title + ': ' #add colon to the end of title
        title = newtitlestring+subtitle #new title is in format title: subtitle.
    else:
        newtitlestring = title + '.' #if subtitle is empty add full stop to the end of title 
        title = newtitlestring #new title is in format title.
    
    p = document.add_paragraph("[%s] " % count) #add a new paragraph to document starting with count value e.g reference [3]
    p.style = document.styles['Normal'] #write paragraph in previously defined font style
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY #justify aligning of paragraph as per required for UoM reports
    p.add_run("%s (%s). " % (author,year)) #add the title and year in brackets to the reference
    p.add_run(title).italic = True #add the title to the paragraph in italics
    if(edition!=''):
        p.add_run(" %s. " % (edition)) #add the edition. publishier location: publisher name. to the paragraph    
    else:
        p.add_run(" ")
    p.add_run(provider).italic = True 
    p.add_run(". [Online]. Available at: %s" % url)
    document.save('references.docx') #save the appending to references.docx
    success = Label(referenceframe,text=('Reference [%s] Successfully Generated' % count)).grid(row=10,column=0,columnspan=2) #create a label at the bottom of the window informing reference [number] added successfully

def organisationreportgen(authorstring,yearstring,titlestring,pubplacestring,publisherstring,urlstring,accessstring,referenceframe):
    global count
    count=count+1 #add one to the count to increase the number in the reference list
    #get the values of each of the inputs of the book window entry fields
    author = authorstring.get()
    year = yearstring.get()
    title = titlestring.get()
    pubplace = pubplacestring.get()
    publisher = publisherstring.get()
    url = urlstring.get()
    access = accessstring.get()
    
    p = document.add_paragraph("[%s] " % count) #add a new paragraph to document starting with count value e.g reference [3]
    p.style = document.styles['Normal'] #write paragraph in previously defined font style
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY #justify aligning of paragraph as per required for UoM reports
    p.add_run("%s (%s). " % (author,year)) #add the title and year in brackets to the reference
    p.add_run(title).italic = True #add the title to the paragraph in italics
    if(publisher!=''):
        p.add_run('. %s: %s.' % (pubplace,publisher))
    else:
        p.add_run('. Available At: %s (Accessed: %s).' % (url,access))
    document.save('references.docx') #save the appending to references.docx
    success = Label(referenceframe,text=('Reference [%s] Successfully Generated' % count)).grid(row=8,column=0,columnspan=2) #create a label at the bottom of the window informing reference [number] added successfully    
    
def selfcitationgen(authorstring,yearstring,titlestring,modulecodestring,moduletitlestring,institutionstring,referenceframe):
    global count
    count=count+1 #add one to the count to increase the number in the reference list
    #get the values of each of the inputs of the book window entry fields
    author = authorstring.get()
    year = yearstring.get()
    title = titlestring.get()
    modulecode = modulecodestring.get()
    moduletitle = moduletitlestring.get()
    institution = institutionstring.get()
    
    p = document.add_paragraph("[%s] " % count) #add a new paragraph to document starting with count value e.g reference [3]
    p.style = document.styles['Normal'] #write paragraph in previously defined font style
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY #justify aligning of paragraph as per required for UoM reports
    p.add_run("%s (%s). '" % (author,year)) #add the title and year in brackets to the reference
    p.add_run(title).italic = True #add the title to the paragraph in italics
    p.add_run("'. %s: " % modulecode)
    p.add_run(moduletitle).italic = True
    p.add_run(', %s. Unpublished assignment.' % institution)
    document.save('references.docx') #save the appending to references.docx
    success = Label(referenceframe,text=('Reference [%s] Successfully Generated' % count)).grid(row=7,column=0,columnspan=2) #create a label at the bottom of the window informing reference [number] added successfully    

def govpubsgen(authorstring,yearstring,titlestring,subtitlestring,pubplacestring,publisherstring,seriesstring,urlstring,accessstring,referenceframe):
    global count
    count=count+1 #add one to the count to increase the number in the reference list
    #get the values of each of the inputs of the book window entry fields
    author = authorstring.get()
    year = yearstring.get()
    title = titlestring.get()
    subtitle = subtitlestring.get()
    pubplace = pubplacestring.get()
    publisher = publisherstring.get()
    series = seriesstring.get()
    url = urlstring.get()
    access = accessstring.get()
    
    if(subtitle!=''): #check if there has been an entry in the subtitle window, if so:
        newsubtitlestring = subtitle + '.' #add full stop to the end of subtitle
        subtitle = newsubtitlestring #overwrite subtitle
        newtitlestring = title + ': ' #add colon to the end of title
        title = newtitlestring+subtitle #new title is in format title: subtitle.
    else:
        newtitlestring = title + '.' #if subtitle is empty add full stop to the end of title 
        title = newtitlestring #new title is in format title.
    
    p = document.add_paragraph("[%s] " % count) #add a new paragraph to document starting with count value e.g reference [3]
    p.style = document.styles['Normal'] #write paragraph in previously defined font style
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY #justify aligning of paragraph as per required for UoM reports
    p.add_run("%s (%s). " % (author,year)) #add the title and year in brackets to the reference
    p.add_run(title).italic = True #add the title to the paragraph in italics
    if(series!=''):
        p.add_run(' %s: %s. (%s).' % (pubplace,publisher,series))
    else:
        p.add_run(' %s: %s.' % (pubplace,publisher))
    if(url!=''):
        p.add_run(' Available at: %s (Accessed: %s).' % (url,access))
    else:
        pass
    document.save('references.docx') #save the appending to references.docx
    success = Label(referenceframe,text=('Reference [%s] Successfully Generated' % count)).grid(row=10,column=0,columnspan=2) #create a label at the bottom of the window informing reference [number] added successfully    

#define bookwindow for reference inputs, called when book source is chosen to overwrite the display
def bookwin():
    global root, homeframe #make the tkinter root and the frame of the home screen globals
    homeframe.destroy() #destroy the homeframe to make the tkinter window blank
    referenceframe = Frame(root) #create a new frame for the references in root
    referenceframe.pack() #pack the new reference frame in the window
    
    #generate label at top of window reminding to enter inputs as shown in example and show in the first row and first and second columns of the grid
    infolabel = Label(referenceframe,text='Enter information for reference. Enter inputs as shown, case and punctuation sensitive.').grid(row=0,column=0,columnspan=2)
    
    #generate label to tell to input author and show example and show in second row and  first column of grid
    authorlabel = Label(referenceframe,text='Author(s) (e.g. Utting, J.)').grid(row=1,column=0)
    #define the string variable the entry will be stored to
    authorstring = StringVar()
    #define the entry window for the text and set it in the first row and second column of the grid
    authorentry = Entry(referenceframe,textvariable=authorstring).grid(row=1,column=1)
    
    #generate label to tell to input year and show example and show in third row and first column of grid
    yearlabel = Label(referenceframe,text='Publication Year (e.g. 2019)').grid(row=2,column=0)
    #define the string variable the entry will be stored to
    yearstring = StringVar()
    #define the entry window for the text and set it in the third row and second column of the grid
    yearstringentry = Entry(referenceframe,textvariable=yearstring).grid(row=2,column=1)
    
    #generate label to tell to input title and show example and show in fourth row and first column of grid
    titlelabel = Label(referenceframe,text='Title Of Book (e.g. Beginner programming)').grid(row=3,column=0)
    #define the string variable the entry will be stored to
    titlestring = StringVar()
    #define the entry window for the text and set it in the fourth row and second column of the grid
    titleentry = Entry(referenceframe,textvariable=titlestring).grid(row=3,column=1)
    
    #generate label to tell to input subtitle and show example and show in fifth row and first column of grid
    subtitlelabel = Label(referenceframe,text='Subtitle (If none leave blank, e.g. testing)').grid(row=4,column=0)
    #define the string variable the entry will be stored to
    subtitlestring = StringVar()
    #define the entry window for the text and set it in the fifth row and second column of the grid
    subtitleentry = Entry(referenceframe,textvariable=subtitlestring).grid(row=4,column=1)
    
    #generate label to tell to input edition and show example and show in sixth row and first column of grid
    editionlabel = Label(referenceframe,text='Edition (e.g. 3rd edn)').grid(row=5,column=0)
    #define the string variable the entry will be stored to
    editionstring = StringVar()
    #define the entry window for the text and set it in the sixth row and second column of the grid
    editionentry = Entry(referenceframe,textvariable=editionstring).grid(row=5,column=1)
    
    #generate label to tell to input publshing place and show example and show in seventh row and first column of grid
    pubplacelabel = Label(referenceframe,text='Place Of Publication (e.g. Manchester)').grid(row=6,column=0)
    #define the string variable the entry will be stored to
    pubplacestring = StringVar()
    #define the entry window for the text and set it in the seventh row and second column of the grid
    pubplaceentry = Entry(referenceframe,textvariable=pubplacestring).grid(row=6,column=1)
    
    #generate label to tell to input publisher and show example and show in eighth row and first column of grid
    publisherlabel = Label(referenceframe,text='Publisher Name (e.g. Pearson)').grid(row=7,column=0)
    #define the string variable the entry will be stored to
    publisherstring = StringVar()
    #define the entry window for the text and set it in the eighth row and second column of the grid
    publisherentry = Entry(referenceframe,textvariable=publisherstring).grid(row=7,column=1)
    
    #generate label to tell to input isbn and show example and show in eighth row and first column of grid
    isbnlabel = Label(referenceframe,text='ISBN-13 Search (e.g. 9780198509509)').grid(row=8,column=0)
    #define the string variable the entry will be stored to
    isbnstring = StringVar()
    #define the entry window for the text and set it in the eighth row and second column of the grid
    isbnentry = Entry(referenceframe,textvariable=isbnstring).grid(row=8,column=1)
    #define button to call the isbnsearch function and show in the eighth row and third column of the grid
    isbnbutton = Button(referenceframe,text='Search',command=lambda: isbnsearch(isbnstring,authorstring,yearstring,titlestring,subtitlestring,editionstring,pubplacestring,publisherstring)).grid(row=8,column=2)
    
    #define a button to generate the reference, it has the bookgen function bound to it so will call it on button press to make a reference from the entered information, show this button in the tenth row and first column, note the ninth row skip is to display the success label in the bookgen field above this
    genrefbutton = Button(referenceframe,text='Generate Reference',command=lambda: bookgen(authorstring,yearstring,titlestring,subtitlestring,editionstring,pubplacestring,publisherstring,referenceframe)).grid(row=10,column=1)
    #define the home button to return from the current reference window back to the sourve selector if desired, calls the home function on button press, show button in tenth row and first column
    homebutton = Button(referenceframe,text='Home',command=lambda: home(referenceframe)).grid(row=10,column=0)
    #define the clear button to empty all fields (actually destroys frame and rebuilds fast enough not to notice in fullscreen), shows this button approximateky midway through (fifth row) and in the third column
    clearbutton = Button(referenceframe,text='Clear',command=lambda: clear(referenceframe,bookwin)).grid(row=4,column=2)
    
    root.update() #used instead of root.mmainloop() to allow for changes using the clear command, updates the window on each loop rather than just loops with the original layout
    return #code should not make it this far but left in for any breakages to try and return to previous function

#define webwindow for reference inputs, works as bookwin
def webwin():
    global root, homeframe
    homeframe.destroy()
    referenceframe = Frame(root)
    referenceframe.pack()
    
    infolabel = Label(referenceframe,text='Enter information for reference. Enter inputs as shown, case and punctuation sensitive.').grid(row=0,column=0,columnspan=2)
    
    authorlabel = Label(referenceframe,text='Author(s) (e.g. Utting, J.)').grid(row=1,column=0)
    authorstring = StringVar()
    authorentry = Entry(referenceframe,textvariable=authorstring).grid(row=1,column=1)
    
    yearlabel = Label(referenceframe,text='Website Publication Year (e.g. 2019)').grid(row=2,column=0)
    yearstring = StringVar()
    yearstringentry = Entry(referenceframe,textvariable=yearstring).grid(row=2,column=1)

    titlelabel = Label(referenceframe,text='Title Of Website (e.g. Beginner programming)').grid(row=3,column=0)
    titlestring = StringVar()
    titleentry = Entry(referenceframe,textvariable=titlestring).grid(row=3,column=1)

    urllabel = Label(referenceframe,text='Website URL (e.g. www.google.co.uk)').grid(row=4,column=0)
    urlstring = StringVar()
    urlentry = Entry(referenceframe,textvariable=urlstring).grid(row=4,column=1)

    accesslabel = Label(referenceframe,text='Access Date (e.g. 20th April 2019)').grid(row=5,column=0)
    accessstring = StringVar()
    accessentry = Entry(referenceframe,textvariable=accessstring).grid(row=5,column=1)
        
    genrefbutton = Button(referenceframe,text='Generate Reference',command=lambda: webgen(authorstring,yearstring,titlestring,urlstring,accessstring,referenceframe)).grid(row=9,column=1)
    homebutton = Button(referenceframe,text='Home',command=lambda: home(referenceframe)).grid(row=9,column=0)
    clearbutton = Button(referenceframe,text='Clear',command=lambda: clear(referenceframe,webwin)).grid(row=3,column=2)
    root.update()
    return

#define journal1window for printed journal reference inputs, works as bookwin
def journal1win():
    global root, homeframe
    homeframe.destroy()
    referenceframe = Frame(root)
    referenceframe.pack()
    
    infolabel = Label(referenceframe,text='Enter information for reference. Enter inputs as shown, case and punctuation sensitive.').grid(row=0,column=0,columnspan=2)
    
    authorlabel = Label(referenceframe,text='Author(s) (e.g. Utting, J.)').grid(row=1,column=0)
    authorstring = StringVar()
    authorentry = Entry(referenceframe,textvariable=authorstring).grid(row=1,column=1)
    
    yearlabel = Label(referenceframe,text='Publication Year (e.g. 2019)').grid(row=2,column=0)
    yearstring = StringVar()
    yearstringentry = Entry(referenceframe,textvariable=yearstring).grid(row=2,column=1)
    
    titlearticlelabel = Label(referenceframe,text='Title Of Article (e.g. Beginner programming)').grid(row=3,column=0)
    titlearticlestring = StringVar()
    titlearticleentry = Entry(referenceframe,textvariable=titlearticlestring).grid(row=3,column=1)

    titlelabel = Label(referenceframe,text='Title Of Journal (e.g. Manchester Programming Journal)').grid(row=4,column=0)
    titlestring = StringVar()
    titleentry = Entry(referenceframe,textvariable=titlestring).grid(row=4,column=1)

    volumelabel = Label(referenceframe,text='Volume Number (e.g. 8)').grid(row=5,column=0)
    volumestring = StringVar()
    volumestringentry = Entry(referenceframe,textvariable=volumestring).grid(row=5,column=1)

    issuelabel = Label(referenceframe,text='Issue/Part Number (e.g. 3)').grid(row=6,column=0)
    issuestring = StringVar()
    issuestringentry = Entry(referenceframe,textvariable=issuestring).grid(row=6,column=1)

    pagelabel = Label(referenceframe,text='Page Range (e.g. 55-59)').grid(row=7,column=0)
    pagestring = StringVar()
    pagestringentry = Entry(referenceframe,textvariable=pagestring).grid(row=7,column=1)

    genrefbutton = Button(referenceframe,text='Generate Reference',command=lambda: journal1gen(authorstring,yearstring,titlearticlestring,titlestring,volumestring,issuestring,pagestring,referenceframe)).grid(row=9,column=1)
    homebutton = Button(referenceframe,text='Home',command=lambda: home(referenceframe)).grid(row=9,column=0)
    clearbutton = Button(referenceframe,text='Clear',command=lambda: clear(referenceframe,journal1win)).grid(row=4,column=2)
    root.update()
    return

#define journal2window for online journal reference inputs, works as bookwin
def journal2win():
    global root, homeframe
    homeframe.destroy()
    referenceframe = Frame(root)
    referenceframe.pack()
    
    infolabel = Label(referenceframe,text='Enter information for reference. Enter inputs as shown, case and punctuation sensitive.').grid(row=0,column=0,columnspan=2)
    
    authorlabel = Label(referenceframe,text='Author(s) (e.g. Utting, J.)').grid(row=1,column=0)
    authorstring = StringVar()
    authorentry = Entry(referenceframe,textvariable=authorstring).grid(row=1,column=1)
    
    yearlabel = Label(referenceframe,text='Publication Year (e.g. 2019)').grid(row=2,column=0)
    yearstring = StringVar()
    yearstringentry = Entry(referenceframe,textvariable=yearstring).grid(row=2,column=1)
    
    titlearticlelabel = Label(referenceframe,text='Title Of Article (e.g. Beginner programming)').grid(row=3,column=0)
    titlearticlestring = StringVar()
    titlearticleentry = Entry(referenceframe,textvariable=titlearticlestring).grid(row=3,column=1)

    titlelabel = Label(referenceframe,text='Title Of Journal (e.g. Manchester Programming Journal)').grid(row=4,column=0)
    titlestring = StringVar()
    titleentry = Entry(referenceframe,textvariable=titlestring).grid(row=4,column=1)

    volumelabel = Label(referenceframe,text='Volume Number (e.g. 8)').grid(row=5,column=0)
    volumestring = StringVar()
    volumestringentry = Entry(referenceframe,textvariable=volumestring).grid(row=5,column=1)

    issuelabel = Label(referenceframe,text='Issue/Part Number (e.g. 3)').grid(row=6,column=0)
    issuestring = StringVar()
    issuestringentry = Entry(referenceframe,textvariable=issuestring).grid(row=6,column=1)

    pagelabel = Label(referenceframe,text='Page Range (e.g. 55-59)').grid(row=7,column=0)
    pagestring = StringVar()
    pagestringentry = Entry(referenceframe,textvariable=pagestring).grid(row=7,column=1)

    urllabel = Label(referenceframe,text='Jounal URL (e.g. www.google.co.uk)').grid(row=8,column=0)
    urlstring = StringVar()
    urlentry = Entry(referenceframe,textvariable=urlstring).grid(row=8,column=1)

    accesslabel = Label(referenceframe,text='Access Date (e.g. 20th April 2019)').grid(row=9,column=0)
    accessstring = StringVar()
    accessentry = Entry(referenceframe,textvariable=accessstring).grid(row=9,column=1)
        
    genrefbutton = Button(referenceframe,text='Generate Reference',command=lambda: journal2gen(authorstring,yearstring,titlearticlestring,titlestring,volumestring,issuestring,pagestring,urlstring,accessstring,referenceframe)).grid(row=11,column=1)
    homebutton = Button(referenceframe,text='Home',command=lambda: home(referenceframe)).grid(row=11,column=0)
    clearbutton = Button(referenceframe,text='Clear',command=lambda: clear(referenceframe,journal2win)).grid(row=5,column=2)
    root.update()
    return

#define lechandoutwindow for reference inputs, works as bookwin
def lechandoutwin():
    global root, homeframe
    homeframe.destroy()
    referenceframe = Frame(root)
    referenceframe.pack()
    
    infolabel = Label(referenceframe,text='Enter information for reference. Enter inputs as shown, case and punctuation sensitive.').grid(row=0,column=0,columnspan=2)
    
    authorlabel = Label(referenceframe,text='Lecturer Name (e.g. Utting, J.)').grid(row=1,column=0)
    authorstring = StringVar()
    authorentry = Entry(referenceframe,textvariable=authorstring).grid(row=1,column=1)
    
    yearlabel = Label(referenceframe,text='Distribution Year (e.g. 2019)').grid(row=2,column=0)
    yearstring = StringVar()
    yearstringentry = Entry(referenceframe,textvariable=yearstring).grid(row=2,column=1)

    titlelabel = Label(referenceframe,text='Title Of Handout (e.g. User interface programming)').grid(row=3,column=0)
    titlestring = StringVar()
    titleentry = Entry(referenceframe,textvariable=titlestring).grid(row=3,column=1)

    codelabel = Label(referenceframe,text='Module Code (e.g. PHYS101)').grid(row=4,column=0)
    codestring = StringVar()
    codeentry = Entry(referenceframe,textvariable=codestring).grid(row=4,column=1)

    modtitlelabel = Label(referenceframe,text='Title Of Module (e.g. Beginner programming)').grid(row=5,column=0)
    modtitlestring = StringVar()
    modtitleentry = Entry(referenceframe,textvariable=modtitlestring).grid(row=5,column=1)

    institutionlabel = Label(referenceframe,text='Insititution Name (e.g. University of Manchester)').grid(row=6,column=0)
    institutionstring = StringVar()
    institutionentry = Entry(referenceframe,textvariable=institutionstring).grid(row=6,column=1)

    genrefbutton = Button(referenceframe,text='Generate Reference',command=lambda: lechandoutgen(authorstring,yearstring,titlestring,codestring,modtitlestring,institutionstring,referenceframe)).grid(row=8,column=1)
    homebutton = Button(referenceframe,text='Home',command=lambda: home(referenceframe)).grid(row=8,column=0)
    clearbutton = Button(referenceframe,text='Clear',command=lambda: clear(referenceframe,lechandoutwin)).grid(row=4,column=2)
    root.update()

#define thesiswindow for reference inputs, works as bookwin
def thesiswin():
    global root, homeframe
    homeframe.destroy()
    referenceframe = Frame(root)
    referenceframe.pack()
    
    infolabel = Label(referenceframe,text='Enter information for reference. Enter inputs as shown, case and punctuation sensitive.').grid(row=0,column=0,columnspan=2)
    
    authorlabel = Label(referenceframe,text='Author(s) (e.g. Utting, J.)').grid(row=1,column=0)
    authorstring = StringVar()
    authorentry = Entry(referenceframe,textvariable=authorstring).grid(row=1,column=1)
    
    yearlabel = Label(referenceframe,text='Submission Year (e.g. 2019)').grid(row=2,column=0)
    yearstring = StringVar()
    yearstringentry = Entry(referenceframe,textvariable=yearstring).grid(row=2,column=1)

    titlelabel = Label(referenceframe,text='Title Of Thesis (e.g. Statistical modelling of neutron scattering)').grid(row=3,column=0)
    titlestring = StringVar()
    titleentry = Entry(referenceframe,textvariable=titlestring).grid(row=3,column=1)
    
    levellabel = Label(referenceframe,text='Level Of Thesis (e.g. PhD)').grid(row=4,column=0)
    levelstring = StringVar()
    levelentry = Entry(referenceframe,textvariable=levelstring).grid(row=4,column=1)
    
    institutionlabel = Label(referenceframe,text='Institution Name (e.g. University of Manchester)').grid(row=5,column=0)
    institutionstring = StringVar()
    institutionentry = Entry(referenceframe,textvariable=institutionstring).grid(row=5,column=1)

    genrefbutton = Button(referenceframe,text='Generate Reference',command=lambda: thesisgen(authorstring,yearstring,titlestring,levelstring,institutionstring,referenceframe)).grid(row=7,column=1)
    homebutton = Button(referenceframe,text='Home',command=lambda: home(referenceframe)).grid(row=7,column=0)
    clearbutton = Button(referenceframe,text='Clear',command=lambda: clear(referenceframe,thesiswin)).grid(row=3,column=2)
    root.update()

#define presentationwindow for reference inputs, works as bookwin
def presentwin():
    global root,homeframe
    homeframe.destroy()
    referenceframe = Frame(root)
    referenceframe.pack()
    
    infolabel = Label(referenceframe,text='Enter information for reference. Enter inputs as shown, case and punctuation sensitive.').grid(row=0,column=0,columnspan=2)
    
    authorlabel = Label(referenceframe,text='Author(s) (e.g. Utting, J.)').grid(row=1,column=0)
    authorstring = StringVar()
    authorentry = Entry(referenceframe,textvariable=authorstring).grid(row=1,column=1)
    
    yearlabel = Label(referenceframe,text='Submission Year (e.g. 2019)').grid(row=2,column=0)
    yearstring = StringVar()
    yearstringentry = Entry(referenceframe,textvariable=yearstring).grid(row=2,column=1)

    titlelabel = Label(referenceframe,text='Title Of Presentation (e.g. Introductory programming)').grid(row=3,column=0)
    titlestring = StringVar()
    titleentry = Entry(referenceframe,textvariable=titlestring).grid(row=3,column=1)
    
    codelabel = Label(referenceframe,text='Module Code (e.g. PHYS101)').grid(row=4,column=0)
    codestring = StringVar()
    codeentry = Entry(referenceframe,textvariable=codestring).grid(row=4,column=1)

    modtitlelabel = Label(referenceframe,text='Title Of Module (e.g. Beginner python)').grid(row=5,column=0)
    modtitlestring = StringVar()
    modtitleentry = Entry(referenceframe,textvariable=modtitlestring).grid(row=5,column=1)

    urllabel = Label(referenceframe,text='Website URL (e.g. www.google.co.uk)').grid(row=6,column=0)
    urlstring = StringVar()
    urlentry = Entry(referenceframe,textvariable=urlstring).grid(row=6,column=1)

    accesslabel = Label(referenceframe,text='Access Date (e.g. 20th April 2019)').grid(row=7,column=0)
    accessstring = StringVar()
    accessentry = Entry(referenceframe,textvariable=accessstring).grid(row=7,column=1)
        
    genrefbutton = Button(referenceframe,text='Generate Reference',command=lambda: presentgen(authorstring,yearstring,titlestring,codestring,modtitlestring,urlstring,accessstring,referenceframe)).grid(row=9,column=1)
    homebutton = Button(referenceframe,text='Home',command=lambda: home(referenceframe)).grid(row=9,column=0)
    clearbutton = Button(referenceframe,text='Clear',command=lambda: clear(referenceframe,presentwin)).grid(row=4,column=2)
    root.update()

#define editedbookwin, same as boow=kwin but for chapter from edited book
def editbookchapterwin():
    global root, homeframe
    homeframe.destroy() 
    referenceframe = Frame(root) 
    referenceframe.pack() 
    
    infolabel = Label(referenceframe,text='Enter information for reference. Enter inputs as shown, case and punctuation sensitive.').grid(row=0,column=0,columnspan=2)
    
    authorlabel = Label(referenceframe,text='Author(s) (e.g. Utting, J.)').grid(row=1,column=0)
    authorstring = StringVar()
    authorentry = Entry(referenceframe,textvariable=authorstring).grid(row=1,column=1)
    
    yearlabel = Label(referenceframe,text='Publication Year (e.g. 2019)').grid(row=2,column=0)
    yearstring = StringVar()
    yearstringentry = Entry(referenceframe,textvariable=yearstring).grid(row=2,column=1)
    
    chapterlabel = Label(referenceframe,text='Title Of Chapter (e.g. Intro to GUI)').grid(row=3,column=0)
    chapterstring = StringVar()
    chapterentry = Entry(referenceframe,textvariable=chapterstring).grid(row=3,column=1)
    
    editorlabel = Label(referenceframe,text='Editor(s) (e.g. Saraili, C.)').grid(row=4,column=0)
    editorstring = StringVar()
    editorentry = Entry(referenceframe,textvariable=editorstring).grid(row=4,column=1)
    
    titlelabel = Label(referenceframe,text='Title Of Book (e.g. Beginner programming)').grid(row=5,column=0)
    titlestring = StringVar()
    titleentry = Entry(referenceframe,textvariable=titlestring).grid(row=5,column=1)
    
    subtitlelabel = Label(referenceframe,text='Subtitle (If none leave blank, e.g. testing)').grid(row=6,column=0)
    subtitlestring = StringVar()
    subtitleentry = Entry(referenceframe,textvariable=subtitlestring).grid(row=6,column=1)
    
    pubplacelabel = Label(referenceframe,text='Place Of Publication (e.g. Manchester)').grid(row=7,column=0)
    pubplacestring = StringVar()
    pubplaceentry = Entry(referenceframe,textvariable=pubplacestring).grid(row=7,column=1)
    
    publisherlabel = Label(referenceframe,text='Publisher Name (e.g. Pearson)').grid(row=8,column=0)
    publisherstring = StringVar()
    publisherentry = Entry(referenceframe,textvariable=publisherstring).grid(row=8,column=1)
    
    pageslabel = Label(referenceframe,text='Page Range (e.g. 36-75)').grid(row=9,column=0)
    pagesstring = StringVar()
    pagesentry = Entry(referenceframe,textvariable=pagesstring).grid(row=9,column=1)
    
    isbnlabel = Label(referenceframe,text='ISBN-13 Search (e.g. 9780198509509)').grid(row=10,column=0)
    isbnstring = StringVar()
    isbnentry = Entry(referenceframe,textvariable=isbnstring).grid(row=10,column=1)
    isbnbutton = Button(referenceframe,text='Search',command=lambda: isbnsearch(isbnstring,authorstring,yearstring,titlestring,subtitlestring,0,pubplacestring,publisherstring)).grid(row=10,column=2)
    
    genrefbutton = Button(referenceframe,text='Generate Reference',command=lambda: editbookchaptergen(authorstring,yearstring,chapterstring,editorstring,titlestring,subtitlestring,pubplacestring,publisherstring,pagesstring,referenceframe)).grid(row=12,column=1)
    homebutton = Button(referenceframe,text='Home',command=lambda: home(referenceframe)).grid(row=12,column=0)
    clearbutton = Button(referenceframe,text='Clear',command=lambda: clear(referenceframe,editbookchapterwin)).grid(row=5,column=2)
    
    root.update() #used instead of root.mmainloop() to allow for changes using the clear command, updates the window on each loop rather than just loops with the original layout
    return #code should not make it this far but left in for any breakages to try and return to previous function    

def editedbookwin():
    global root, homeframe
    homeframe.destroy() 
    referenceframe = Frame(root) 
    referenceframe.pack() 
    
    infolabel = Label(referenceframe,text='Enter information for reference. Enter inputs as shown, case and punctuation sensitive.').grid(row=0,column=0,columnspan=2)
    
    authorlabel = Label(referenceframe,text='Author(s) (e.g. Utting, J.)').grid(row=1,column=0)
    authorstring = StringVar()
    authorentry = Entry(referenceframe,textvariable=authorstring).grid(row=1,column=1)
    
    yearlabel = Label(referenceframe,text='Publication Year (e.g. 2019)').grid(row=2,column=0)
    yearstring = StringVar()
    yearstringentry = Entry(referenceframe,textvariable=yearstring).grid(row=2,column=1)
    
    titlelabel = Label(referenceframe,text='Title Of Book (e.g. Beginner programming)').grid(row=3,column=0)
    titlestring = StringVar()
    titleentry = Entry(referenceframe,textvariable=titlestring).grid(row=3,column=1)
    
    subtitlelabel = Label(referenceframe,text='Subtitle (If none leave blank, e.g. testing)').grid(row=4,column=0)
    subtitlestring = StringVar()
    subtitleentry = Entry(referenceframe,textvariable=subtitlestring).grid(row=4,column=1)
    
    editionlabel = Label(referenceframe,text='Edition (If first leave blank, e.g. 3rd edn)').grid(row=5,column=0)
    editionstring = StringVar()
    editionentry = Entry(referenceframe,textvariable=editionstring).grid(row=5,column=1)
    
    pubplacelabel = Label(referenceframe,text='Place Of Publication (e.g. Manchester)').grid(row=6,column=0)
    pubplacestring = StringVar()
    pubplaceentry = Entry(referenceframe,textvariable=pubplacestring).grid(row=6,column=1)
    
    publisherlabel = Label(referenceframe,text='Publisher Name (e.g. Pearson)').grid(row=7,column=0)
    publisherstring = StringVar()
    publisherentry = Entry(referenceframe,textvariable=publisherstring).grid(row=7,column=1)
    
    isbnlabel = Label(referenceframe,text='ISBN-13 Search (e.g. 9780198509509)').grid(row=8,column=0)
    isbnstring = StringVar()
    isbnentry = Entry(referenceframe,textvariable=isbnstring).grid(row=8,column=1)
    isbnbutton = Button(referenceframe,text='Search',command=lambda: isbnsearch(isbnstring,authorstring,yearstring,titlestring,subtitlestring,0,pubplacestring,publisherstring)).grid(row=8,column=2)
    
    genrefbutton = Button(referenceframe,text='Generate Reference',command=lambda: editedbookgen(authorstring,yearstring,titlestring,subtitlestring,editionstring,pubplacestring,publisherstring,referenceframe)).grid(row=10,column=1)
    homebutton = Button(referenceframe,text='Home',command=lambda: home(referenceframe)).grid(row=10,column=0)
    clearbutton = Button(referenceframe,text='Clear',command=lambda: clear(referenceframe,editedbookwin)).grid(row=4,column=2)
    
    root.update() #used instead of root.mmainloop() to allow for changes using the clear command, updates the window on each loop rather than just loops with the original layout
    return #code should not make it this far but left in for any breakages to try and return to previous function    

def softwarewin():
    global root, homeframe
    homeframe.destroy() 
    referenceframe = Frame(root) 
    referenceframe.pack() 
    
    infolabel = Label(referenceframe,text='Enter information for reference. Enter inputs as shown, case and punctuation sensitive.').grid(row=0,column=0,columnspan=2)
    
    authorlabel = Label(referenceframe,text='Author(s) (e.g. Utting, J.)').grid(row=1,column=0)
    authorstring = StringVar()
    authorentry = Entry(referenceframe,textvariable=authorstring).grid(row=1,column=1)
    
    yearlabel = Label(referenceframe,text='Release Year (e.g. 2019)').grid(row=2,column=0)
    yearstring = StringVar()
    yearstringentry = Entry(referenceframe,textvariable=yearstring).grid(row=2,column=1)
    
    titlelabel = Label(referenceframe,text='Title Of Software (e.g. Harvard Reference Generator)').grid(row=3,column=0)
    titlestring = StringVar()
    titleentry = Entry(referenceframe,textvariable=titlestring).grid(row=3,column=1)
    
    versionlabel = Label(referenceframe,text='Version Of Software (e.g. Version 2)').grid(row=4,column=0)
    versionstring = StringVar()
    versionentry = Entry(referenceframe,textvariable=versionstring).grid(row=4,column=1)
    
    urllabel = Label(referenceframe,text='URL (e.g. www.google.co.uk)').grid(row=5,column=0)
    urlstring = StringVar()
    urlentry = Entry(referenceframe,textvariable=urlstring).grid(row=5,column=1)
    
    accesslabel = Label(referenceframe,text='Access Date (e.g. 20th April 2019)').grid(row=6,column=0)
    accessstring = StringVar()
    accessentry = Entry(referenceframe,textvariable=accessstring).grid(row=6,column=1)

    genrefbutton = Button(referenceframe,text='Generate Reference',command=lambda: softwaregen(authorstring,yearstring,titlestring,versionstring,urlstring,accessstring,referenceframe)).grid(row=8,column=1)
    homebutton = Button(referenceframe,text='Home',command=lambda: home(referenceframe)).grid(row=8,column=0)
    clearbutton = Button(referenceframe,text='Clear',command=lambda: clear(referenceframe,softwarewin)).grid(row=4,column=2)
    root.update()
    return

def illustrationwin():
    global root, homeframe
    homeframe.destroy() 
    referenceframe = Frame(root) 
    referenceframe.pack() 
    
    infolabel = Label(referenceframe,text='Enter information for reference. Enter inputs as shown, case and punctuation sensitive.').grid(row=0,column=0,columnspan=2)
    
    authorlabel = Label(referenceframe,text='Author(s) (e.g. Utting, J.)').grid(row=1,column=0)
    authorstring = StringVar()
    authorentry = Entry(referenceframe,textvariable=authorstring).grid(row=1,column=1)
    
    yearlabel = Label(referenceframe,text='Publication Year (e.g. 2019)').grid(row=2,column=0)
    yearstring = StringVar()
    yearstringentry = Entry(referenceframe,textvariable=yearstring).grid(row=2,column=1)
    
    titlelabel = Label(referenceframe,text='Title Of Image (e.g. Beginner programming)').grid(row=3,column=0)
    titlestring = StringVar()
    titleentry = Entry(referenceframe,textvariable=titlestring).grid(row=3,column=1)
    
    subtitlelabel = Label(referenceframe,text='Subtitle (If none leave blank, e.g. testing)').grid(row=4,column=0)
    subtitlestring = StringVar()
    subtitleentry = Entry(referenceframe,textvariable=subtitlestring).grid(row=4,column=1)
    
    pubplacelabel = Label(referenceframe,text='Place Of Publication (If in book else leave blank, e.g. Manchester)').grid(row=5,column=0)
    pubplacestring = StringVar()
    pubplaceentry = Entry(referenceframe,textvariable=pubplacestring).grid(row=5,column=1)
    
    publisherlabel = Label(referenceframe,text='Publisher Name (If in book else leave blank, e.g. Pearson)').grid(row=6,column=0)
    publisherstring = StringVar()
    publisherentry = Entry(referenceframe,textvariable=publisherstring).grid(row=6,column=1)
    
    pagelabel = Label(referenceframe,text='Page Number (If in book else leave blank, e.g. 23)').grid(row=7,column=0)
    pagestring = StringVar()
    pageentry = Entry(referenceframe,textvariable=pagestring).grid(row=7,column=1)
    
    urllabel = Label(referenceframe,text='URL (If on website else leave blank, e.g. www.google.co.uk)').grid(row=8,column=0)
    urlstring = StringVar()
    urlentry = Entry(referenceframe,textvariable=urlstring).grid(row=8,column=1)
    
    accesslabel = Label(referenceframe,text='Access Date (e.g. 20th April 2019)').grid(row=9,column=0)
    accessstring = StringVar()
    accessentry = Entry(referenceframe,textvariable=accessstring).grid(row=9,column=1)
    
    typelabel = Label(referenceframe,text='Type Of Medium (e.g. illustration/diagram/table)').grid(row=10,column=0)
    typestring = StringVar()
    typeentry = Entry(referenceframe,textvariable=typestring).grid(row=10,column=1)
    
    genrefbutton = Button(referenceframe,text='Generate Reference',command=lambda: illustrationgen(authorstring,yearstring,titlestring,subtitlestring,pubplacestring,publisherstring,pagestring,urlstring,accessstring,typestring,referenceframe)).grid(row=12,column=1)
    homebutton = Button(referenceframe,text='Home',command=lambda: home(referenceframe)).grid(row=12,column=0)
    clearbutton = Button(referenceframe,text='Clear',command=lambda: clear(referenceframe,illustrationwin)).grid(row=5,column=2)
    root.update()
    return

def ebookwin():
    global root, homeframe
    homeframe.destroy() 
    referenceframe = Frame(root) 
    referenceframe.pack() 
    
    infolabel = Label(referenceframe,text='Enter information for reference. Enter inputs as shown, case and punctuation sensitive.').grid(row=0,column=0,columnspan=2)
    
    authorlabel = Label(referenceframe,text='Author(s) (e.g. Utting, J.)').grid(row=1,column=0)
    authorstring = StringVar()
    authorentry = Entry(referenceframe,textvariable=authorstring).grid(row=1,column=1)
    
    yearlabel = Label(referenceframe,text='Publication Year (e.g. 2019)').grid(row=2,column=0)
    yearstring = StringVar()
    yearstringentry = Entry(referenceframe,textvariable=yearstring).grid(row=2,column=1)
    
    titlelabel = Label(referenceframe,text='Title Of E-Book (e.g. Beginner programming)').grid(row=3,column=0)
    titlestring = StringVar()
    titleentry = Entry(referenceframe,textvariable=titlestring).grid(row=3,column=1)
    
    subtitlelabel = Label(referenceframe,text='Subtitle (If none leave blank, e.g. testing)').grid(row=4,column=0)
    subtitlestring = StringVar()
    subtitleentry = Entry(referenceframe,textvariable=subtitlestring).grid(row=4,column=1)
    
    editionlabel = Label(referenceframe,text='Edition (If first leave blank, e.g. 3rd edn)').grid(row=5,column=0)
    editionstring = StringVar()
    editionentry = Entry(referenceframe,textvariable=editionstring).grid(row=5,column=1)
    
    providerlabel = Label(referenceframe,text='E-Book Provider (e.g. Knovel)').grid(row=6,column=0)
    providerstring = StringVar()
    providerentry = Entry(referenceframe,textvariable=providerstring).grid(row=6,column=1)
    
    urllabel = Label(referenceframe,text='Host URL (e.g. www.google.co.uk)').grid(row=7,column=0)
    urlstring = StringVar()
    urlentry = Entry(referenceframe,textvariable=urlstring).grid(row=7,column=1)
    
    genrefbutton = Button(referenceframe,text='Generate Reference',command=lambda: ebookgen(authorstring,yearstring,titlestring,subtitlestring,editionstring,providerstring,urlstring,referenceframe)).grid(row=9,column=1)
    homebutton = Button(referenceframe,text='Home',command=lambda: home(referenceframe)).grid(row=9,column=0)
    clearbutton = Button(referenceframe,text='Clear',command=lambda: clear(referenceframe,ebookwin)).grid(row=4,column=2)
    
    root.update() #used instead of root.mmainloop() to allow for changes using the clear command, updates the window on each loop rather than just loops with the original layout
    return #code should not make it this far but left in for any breakages to try and return to previous function    

def organisationreportwin():
    global root, homeframe
    homeframe.destroy() 
    referenceframe = Frame(root) 
    referenceframe.pack() 
    
    infolabel = Label(referenceframe,text='Enter information for reference. Enter inputs as shown, case and punctuation sensitive.').grid(row=0,column=0,columnspan=2)
    
    authorlabel = Label(referenceframe,text='Organisation (e.g. British Programming Council)').grid(row=1,column=0)
    authorstring = StringVar()
    authorentry = Entry(referenceframe,textvariable=authorstring).grid(row=1,column=1)
    
    yearlabel = Label(referenceframe,text='Publication Year (e.g. 2019)').grid(row=2,column=0)
    yearstring = StringVar()
    yearstringentry = Entry(referenceframe,textvariable=yearstring).grid(row=2,column=1)
    
    titlelabel = Label(referenceframe,text='Title Of Report (e.g. Beginner programming)').grid(row=3,column=0)
    titlestring = StringVar()
    titleentry = Entry(referenceframe,textvariable=titlestring).grid(row=3,column=1)
    
    pubplacelabel = Label(referenceframe,text='Place Of Publication (If published else leave blank, e.g. Manchester)').grid(row=4,column=0)
    pubplacestring = StringVar()
    pubplaceentry = Entry(referenceframe,textvariable=pubplacestring).grid(row=4,column=1)
    
    publisherlabel = Label(referenceframe,text='Publisher Name (If published else leave blank, e.g. Pearson)').grid(row=5,column=0)
    publisherstring = StringVar()
    publisherentry = Entry(referenceframe,textvariable=publisherstring).grid(row=5,column=1)
    
    urllabel = Label(referenceframe,text='URL (If on website else leave blank, e.g. www.google.co.uk)').grid(row=6,column=0)
    urlstring = StringVar()
    urlentry = Entry(referenceframe,textvariable=urlstring).grid(row=6,column=1)
    
    accesslabel = Label(referenceframe,text='Access Date (If on website else leave blank, e.g. 20th April 2019)').grid(row=7,column=0)
    accessstring = StringVar()
    accessentry = Entry(referenceframe,textvariable=accessstring).grid(row=7,column=1)
    
    genrefbutton = Button(referenceframe,text='Generate Reference',command=lambda: organisationreportgen(authorstring,yearstring,titlestring,pubplacestring,publisherstring,urlstring,accessstring,referenceframe)).grid(row=9,column=1)
    homebutton = Button(referenceframe,text='Home',command=lambda: home(referenceframe)).grid(row=9,column=0)
    clearbutton = Button(referenceframe,text='Clear',command=lambda: clear(referenceframe,organisationreportwin)).grid(row=4,column=2)
    
    root.update() #used instead of root.mmainloop() to allow for changes using the clear command, updates the window on each loop rather than just loops with the original layout
    return #code should not make it this far but left in for any breakages to try and return to previous function   

def selfcitationwin():
    global root, homeframe
    homeframe.destroy() 
    referenceframe = Frame(root) 
    referenceframe.pack() 
    
    infolabel = Label(referenceframe,text='Enter information for reference. Enter inputs as shown, case and punctuation sensitive.').grid(row=0,column=0,columnspan=2)
    
    authorlabel = Label(referenceframe,text='Student Name (e.g. Utting, Joshua)').grid(row=1,column=0)
    authorstring = StringVar()
    authorentry = Entry(referenceframe,textvariable=authorstring).grid(row=1,column=1)
    
    yearlabel = Label(referenceframe,text='Submission Year (e.g. 2019)').grid(row=2,column=0)
    yearstring = StringVar()
    yearstringentry = Entry(referenceframe,textvariable=yearstring).grid(row=2,column=1)
    
    titlelabel = Label(referenceframe,text='Title Of Essay/Assignment (e.g. Beginner programming)').grid(row=3,column=0)
    titlestring = StringVar()
    titleentry = Entry(referenceframe,textvariable=titlestring).grid(row=3,column=1)

    modulecodelabel = Label(referenceframe,text='Module Code (e.g. PHYS10101)').grid(row=4,column=0)
    modulecodestring = StringVar()
    modulecodeentry = Entry(referenceframe,textvariable=modulecodestring).grid(row=4,column=1)

    moduletitlelabel = Label(referenceframe,text='Module Title (e.g. Introduction to Programming)').grid(row=5,column=0)
    moduletitlestring = StringVar()
    moduletitleentry = Entry(referenceframe,textvariable=moduletitlestring).grid(row=5,column=1)
    
    institutionlabel = Label(referenceframe,text='Institution (e.g. University of Manchester)').grid(row=6,column=0)
    institutionstring = StringVar()
    institutionentry = Entry(referenceframe,textvariable=institutionstring).grid(row=6,column=1)
    
    genrefbutton = Button(referenceframe,text='Generate Reference',command=lambda: selfcitationgen(authorstring,yearstring,titlestring,modulecodestring,moduletitlestring,institutionstring,referenceframe)).grid(row=8,column=1)
    homebutton = Button(referenceframe,text='Home',command=lambda: home(referenceframe)).grid(row=8,column=0)
    clearbutton = Button(referenceframe,text='Clear',command=lambda: clear(referenceframe,selfcitationwin)).grid(row=3,column=2)

def govpubswin():
    global root, homeframe
    homeframe.destroy() 
    referenceframe = Frame(root) 
    referenceframe.pack() 
    
    infolabel = Label(referenceframe,text='Enter information for reference. Enter inputs as shown, case and punctuation sensitive.').grid(row=0,column=0,columnspan=2)
    
    authorlabel = Label(referenceframe,text='Department Name (e.g. Department of Technology)').grid(row=1,column=0)
    authorstring = StringVar()
    authorentry = Entry(referenceframe,textvariable=authorstring).grid(row=1,column=1)
    
    yearlabel = Label(referenceframe,text='Publication Year (e.g. 2019)').grid(row=2,column=0)
    yearstring = StringVar()
    yearstringentry = Entry(referenceframe,textvariable=yearstring).grid(row=2,column=1)
    
    titlelabel = Label(referenceframe,text='Title Of Publication (e.g. Beginner programming guidelines)').grid(row=3,column=0)
    titlestring = StringVar()
    titleentry = Entry(referenceframe,textvariable=titlestring).grid(row=3,column=1)
    
    subtitlelabel = Label(referenceframe,text='Subtitle (If none leave blank, e.g. testing)').grid(row=4,column=0)
    subtitlestring = StringVar()
    subtitleentry = Entry(referenceframe,textvariable=subtitlestring).grid(row=4,column=1)
    
    pubplacelabel = Label(referenceframe,text='Place Of Publication (e.g. Manchester)').grid(row=5,column=0)
    pubplacestring = StringVar()
    pubplaceentry = Entry(referenceframe,textvariable=pubplacestring).grid(row=5,column=1)
    
    publisherlabel = Label(referenceframe,text='Publisher Name (e.g. Pearson)').grid(row=6,column=0)
    publisherstring = StringVar()
    publisherentry = Entry(referenceframe,textvariable=publisherstring).grid(row=6,column=1)
    
    serieslabel = Label(referenceframe,text='Series (If none leave blank e.g. No. 959)').grid(row=7,column=0)
    seriesstring = StringVar()
    seriesentry = Entry(referenceframe,textvariable=seriesstring).grid(row=7,column=1)
    
    urllabel = Label(referenceframe,text='URL (If on website else leave blank, e.g. www.google.co.uk)').grid(row=8,column=0)
    urlstring = StringVar()
    urlentry = Entry(referenceframe,textvariable=urlstring).grid(row=8,column=1)
    
    accesslabel = Label(referenceframe,text='Access Date (If on website else leave blank, e.g. 20th April 2019)').grid(row=9,column=0)
    accessstring = StringVar()
    accessentry = Entry(referenceframe,textvariable=accessstring).grid(row=9,column=1)
    
    genrefbutton = Button(referenceframe,text='Generate Reference',command=lambda: govpubsgen(authorstring,yearstring,titlestring,subtitlestring,pubplacestring,publisherstring,seriesstring,urlstring,accessstring,referenceframe)).grid(row=11,column=1)
    homebutton = Button(referenceframe,text='Home',command=lambda: home(referenceframe)).grid(row=11,column=0)
    clearbutton = Button(referenceframe,text='Clear',command=lambda: clear(referenceframe,govpubswin)).grid(row=4,column=2)

def customhelp():
    helpwin = Tk()
    helpwin.title('Help')
    helpwin.attributes("-topmost", True) 
    infolabel1 = Label(helpwin,text='This mode is used for creating your own reference style if it is non standard or just not included so far.').grid(row=0,column=0)
    infolabel2 = Label(helpwin,text='For each element to add there is four options, punctuation is for the punctuation following or around that element. If it is a single piece it will follow the element and format e.g. string. for the full stop option.\nIf there are two pieces or more it will place one before and one after the element and format e.g. (string). for the brackets and full stop option.').grid(row=1,column=0)
    infolabel3 = Label(helpwin,text='Click add element for as many elements are required and type the information such as author in the entry field. Choose the appropriate punctuation to run on.').grid(row=2,column=0)
    infolabel4 = Label(helpwin,text='There are two checkboxes on each element, one for italics and one for bold, check these if the appropriate formatting is required for the element and choose both if both are required.').grid(row=3,column=0)
    infolabel4 = Label(helpwin,text='An example of how to sue this function can be seen below, this combination of inputs will create the same as the book generation function to help you understand how this works :').grid(row=4,column=0)
    
    image = Image.open('CustomExample.png')
    f = Figure(figsize=(13,6))
    a = f.add_subplot(111)
    a.imshow(image)
    a.axis('off')
    canvas = FigureCanvasTkAgg(f,master=helpwin)
    canvas.get_tk_widget().grid(row=5,column=0)
    canvas._tkcanvas.grid(row=5,column=0)
    
    helpwin.mainloop()

def customgen(puncvars,stringvariables,italvars,boldvars,frame):
    global count
    count+=1
    for i in reversed(range(0,len(stringvariables))):
        if(stringvariables[i].get()==''):
            stringvariables.pop(i)
        else:
            stringvariables[i] = stringvariables[i].get()
    print(stringvariables)
    totalfield = []
    for i in range(0,len(stringvariables)):
        puncvar = puncvars[i].get()
        italvar = italvars[i].get()
        boldvar = boldvars[i].get()
        totalfield.append([puncvar,stringvariables[i],italvar,boldvar])
        print(totalfield[i])
    
    p = document.add_paragraph("[%s] " % count) #add a new paragraph to document starting with count value e.g reference [3]
    p.style = document.styles['Normal'] #write paragraph in previously defined font style
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY #justify aligning of paragraph as per required for UoM reports
    for i in range(0,len(totalfield)):
        if(totalfield[i][0]=='.'):
            if(totalfield[i][2]==1 and totalfield[i][3]==1):
                runner = p.add_run('%s. ' % totalfield[i][1])
                runner.bold=True
                runner.italic=True
            elif(totalfield[i][2]==1 and totalfield[i][3]!=1):
                p.add_run('%s. ' % totalfield[i][1]).italic = True
            elif(totalfield[i][3]==1 and totalfield[i][2]!=1):
                p.add_run('%s. ' % totalfield[i][1]).bold = True
            else:
                p.add_run('%s. ' % totalfield[i][1])
        elif(totalfield[i][0]==','):
            if(totalfield[i][2]==1 and totalfield[i][3]==1):
                runner = p.add_run('%s, ' % totalfield[i][1])
                runner.bold=True
                runner.italic=True
            elif(totalfield[i][2]==1 and totalfield[i][3]!=1):
                p.add_run('%s, ' % totalfield[i][1]).italic = True
            elif(totalfield[i][3]==1 and totalfield[i][2]!=1):
                p.add_run('%s, ' % totalfield[i][1]).bold = True
            else:
                p.add_run('%s, ' % totalfield[i][1])
        elif(totalfield[i][0]=='( )'):
            if(totalfield[i][2]==1 and totalfield[i][3]==1):
                runner = p.add_run('(%s) ' % totalfield[i][1])
                runner.bold=True
                runner.italic=True
            elif(totalfield[i][2]==1 and totalfield[i][3]!=1):
                p.add_run('(%s) ' % totalfield[i][1]).italic = True
            elif(totalfield[i][3]==1 and totalfield[i][2]!=1):
                p.add_run('(%s) ' % totalfield[i][1]).bold = True
            else:
                p.add_run('(%s) ' % totalfield[i][1])
        elif(totalfield[i][0]=='( ).'):
            if(totalfield[i][2]==1 and totalfield[i][3]==1):
                runner = p.add_run('(%s). ' % totalfield[i][1])
                runner.bold=True
                runner.italic=True
            elif(totalfield[i][2]==1 and totalfield[i][3]!=1):
                p.add_run('(%s). ' % totalfield[i][1]).italic = True
            elif(totalfield[i][3]==1 and totalfield[i][2]!=1):
                p.add_run('(%s). ' % totalfield[i][1]).bold = True
            else:
                p.add_run('(%s). ' % totalfield[i][1])
        elif(totalfield[i][0]=='( ),'):
            if(totalfield[i][2]==1 and totalfield[i][3]==1):
                runner = p.add_run('(%s), ' % totalfield[i][1])
                runner.bold=True
                runner.italic=True
            elif(totalfield[i][2]==1 and totalfield[i][3]!=1):
                p.add_run('(%s), ' % totalfield[i][1]).italic = True
            elif(totalfield[i][3]==1 and totalfield[i][2]!=1):
                p.add_run('(%s), ' % totalfield[i][1]).bold = True
            else:
                p.add_run('(%s), ' % totalfield[i][1])
        elif(totalfield[i][0]==':'):
            if(totalfield[i][2]==1 and totalfield[i][3]==1):
                runner = p.add_run('%s: ' % totalfield[i][1])
                runner.bold=True
                runner.italic=True
            if(totalfield[i][2]==1 and totalfield[i][3]!=1):
                p.add_run('%s: ' % totalfield[i][1]).italic = True
            elif(totalfield[i][3]==1 and totalfield[i][2]!=1):
                p.add_run('%s: ' % totalfield[i][1]).bold = True
            else:
                p.add_run('%s: ' % totalfield[i][1])
        elif(totalfield[i][0]=="' '"):
            if(totalfield[i][2]==1 and totalfield[i][3]==1):
                runner = p.add_run("'%s' " % totalfield[i][1])
                runner.bold=True
                runner.italic=True
            elif(totalfield[i][2]==1 and totalfield[i][3]!=1):
                p.add_run("'%s' " % totalfield[i][1]).italic = True
            elif(totalfield[i][3]==1 and totalfield[i][2]!=1):
                p.add_run("'%s' " % totalfield[i][1]).bold = True
            else:
                p.add_run("'%s' " % totalfield[i][1])
        elif(totalfield[i][0]=="' '."):
            if(totalfield[i][2]==1 and totalfield[i][3]==1):
                runner = p.add_run("'%s'. " % totalfield[i][1])
                runner.bold=True
                runner.italic=True
            elif(totalfield[i][2]==1 and totalfield[i][3]!=1):
                p.add_run("'%s'. " % totalfield[i][1]).italic = True
            elif(totalfield[i][3]==1 and totalfield[i][2]!=1):
                p.add_run("'%s'. " % totalfield[i][1]).bold = True
            else:
                p.add_run("'%s'. " % totalfield[i][1])
        elif(totalfield[i][0]=="' ',"):
            if(totalfield[i][2]==1 and totalfield[i][3]==1):
                runner = p.add_run("'%s', " % totalfield[i][1])
                runner.bold=True
                runner.italic=True
            elif(totalfield[i][2]==1 and totalfield[i][3]!=1):
                p.add_run("'%s', " % totalfield[i][1]).italic = True
            elif(totalfield[i][3]==1 and totalfield[i][2]!=1):
                p.add_run("'%s', " % totalfield[i][1]).bold = True
            else:
                p.add_run("'%s', " % totalfield[i][1])
        elif(totalfield[i][0]=='" "'):
            if(totalfield[i][2]==1 and totalfield[i][3]==1):
                runner = p.add_run('"%s" ' % totalfield[i][1])
                runner.bold=True
                runner.italic=True
            elif(totalfield[i][2]==1 and totalfield[i][3]!=1):
                p.add_run('"%s" ' % totalfield[i][1]).italic = True
            elif(totalfield[i][3]==1 and totalfield[i][2]!=1):
                p.add_run('"%s" ' % totalfield[i][1]).bold = True
            else:
                p.add_run('"%s" ' % totalfield[i][1])
        elif(totalfield[i][0]=='" ",'):
            if(totalfield[i][2]==1 and totalfield[i][3]==1):
                runner = p.add_run('"%s", ' % totalfield[i][1])
                runner.bold=True
                runner.italic=True
            elif(totalfield[i][2]==1 and totalfield[i][3]!=1):
                p.add_run('"%s", ' % totalfield[i][1]).italic = True
            elif(totalfield[i][3]==1 and totalfield[i][2]!=1):
                p.add_run('"%s", ' % totalfield[i][1]).bold = True
            else:
                p.add_run('"%s", ' % totalfield[i][1])
        elif(totalfield[i][0]=='" ".'):
            if(totalfield[i][2]==1 and totalfield[i][3]==1):
                runner = p.add_run('"%s". ' % totalfield[i][1])
                runner.bold=True
                runner.italic=True
            elif(totalfield[i][2]==1 and totalfield[i][3]!=1):
                p.add_run('"%s". ' % totalfield[i][1]).italic = True
            elif(totalfield[i][3]==1 and totalfield[i][2]!=1):
                p.add_run('"%s". ' % totalfield[i][1]).bold = True
            else:
                p.add_run('"%s". ' % totalfield[i][1])
        elif(totalfield[i][0]=='None'):
            if(totalfield[i][2]==1 and totalfield[i][3]==1):
                runner = p.add_run('%s ' % totalfield[i][1])
                runner.bold=True
                runner.italic=True
            elif(totalfield[i][2]==1 and totalfield[i][3]!=1):
                p.add_run('%s ' % totalfield[i][1]).italic = True
            elif(totalfield[i][3]==1 and totalfield[i][2]!=1):
                p.add_run('%s ' % totalfield[i][1]).bold = True
            else:
                p.add_run('%s ' % totalfield[i][1])
    document.save('references.docx') #save the appending to references.docx
    success = Label(frame,text=('Reference [%s] Successfully Generated' % count)).grid(row=23,column=0,columnspan=4) #create a label at the bottom of the window informing reference [number] added successfully
    return

def addelement(frame,puncvariables,stringvariables,italints,boldints):
    global elementscount
    if(elementscount<22):
        puncchoices = ['.',',','( )','( ).','( ),',':',"' '","' '.","' ',",'" "','" ",','" ".','None']
        puncmenu = OptionMenu(frame,puncvariables[(elementscount-2)],*puncchoices)
        puncmenu.config(relief='solid')
        puncmenu["highlightthickness"]=0
        puncmenu.grid(row=elementscount,column=0,sticky='nsew')
        newentry = Entry(frame,textvariable=stringvariables[(elementscount-2)],relief='solid').grid(row=elementscount,column=1,sticky='nsew')
        italicbutton = Checkbutton(frame,variable=italints[(elementscount-2)],relief='solid').grid(row=elementscount,column=2,sticky='nsew')
        boldbutton = Checkbutton(frame,variable=boldints[(elementscount-2)],relief='solid').grid(row=elementscount,column=3,sticky='nsew')
        elementscount+=1
    else:
        label = Label(frame,text='Sorry you can only use up to 20 elements in custom mode.',relief='solid').grid(row=elementscount,column=0,columnspan=4,sticky='nsew')
    return

def customwin():
    global root, homeframe, elementscount
    homeframe.destroy() 
    referenceframe = Frame(root) 
    referenceframe.pack() 
    varslist = []
    punclist = []
    itallist = []
    boldlist = []
    
    for i in range(0,20):
        textvar = StringVar()
        varslist.append(textvar)
        puncvar = StringVar()
        puncvar.set('.')
        punclist.append(puncvar)
        italicvar = IntVar()
        itallist.append(italicvar)
        boldvar = IntVar()
        boldlist.append(boldvar)
         
    infolabel = Label(referenceframe,text='This mode is for advanced users only, use it to build your own reference type if your source does not match an existing listed one.\nCheck with library services what elements you need to build your own.').grid(row=0,column=0,columnspan=4)
    punclabel = Label(referenceframe,text='Punctuation:',relief='solid').grid(row=1,column=0,sticky='nsew')
    elementlabel = Label(referenceframe,text='Element Entry:',relief='solid').grid(row=1,column=1,sticky='nsew')
    italiclabel = Label(referenceframe,text='Italicised?',relief='solid').grid(row=1,column=2,sticky='nsew')
    boldlabel = Label(referenceframe,text='Bold?',relief='solid').grid(row=1,column=3,sticky='nsew')
    addfieldbutton = Button(referenceframe,text='Add Element',command=lambda: addelement(referenceframe,punclist,varslist,itallist,boldlist)).grid(row=24,column=0,sticky='nsew')
    genrefbutton = Button(referenceframe,text='Generate Reference',command=lambda: customgen(punclist,varslist,itallist,boldlist,referenceframe)).grid(row=24,column=1,sticky='nsew')
    helpbutton = Button(referenceframe,text='Help',command=customhelp).grid(row=24,column=2,sticky='nsew')
    homebutton = Button(referenceframe,text='Home',command=lambda: home(referenceframe)).grid(row=24,column=3,sticky='nsew')
    addelement(referenceframe,punclist,varslist,itallist,boldlist)
    root.update()    
   
#define quitwindow function to be called by quit buttons
def quitwindow():
    global root #make access to root global
    root.destroy() #destroy root to end program

#define home function to be called by home buttons
def home(frame):
    global elementscount
    elementscount=2
    frame.destroy() #destroy imported frame, in this case the referenceframe of the previously chosen reference
    main() #call main again to show source selector

#define clear funciton to clear current entries if the same source wants to be used again
def clear(frame,functionreturn):
    global root #make access to root global
    frame.destroy() #destroy the current frame that is given in the call of this function
    functionreturn() #call the function window to redraw the input fields of the source this function was called from

#define main function for source selection window
def main():
    global root,homeframe,reftypevar,countvar #make access to root, homeframe, referency type variable and the count variable global
    
    homeframe = Frame(root) #create a frame called homeframe in root
    homeframe.pack() #pack the frame in the root window
    
    reftypevar=StringVar() #define reference type variable as string
    choices = ['Book','Chapter From Edited Book','E-Book','Edited Book','Government/Corporate Publications','Illustration','Journal (electronic)','Journal (printed)','Lecturer Handout','Presentation','Report From Organisation','Self-Citation','Software','Thesis','Website','---------------------------------------','Custom'] #define drop down menu choices for sources
    reftypevar.set('Book') #set default value to book
    
    droplabel = Label(homeframe,text='Choose Source').grid(row=0,column=0) #define label for drop menu and pack it to the first row and first column
    dropmenu = OptionMenu(homeframe,reftypevar,*choices).grid(row=0,column=1) #define drop down menu with the previously defined choices and pack to first row and second column
    #define the confirm button which calls the function to find out what selection was made and calls the appropriate window to generate, pack it to the first row and third column
    refgenbutton = Button(homeframe,text='Confirm',command=getchoice).grid(row=0,column=2)
    #define the quitbutton which calls the quitwindow function on buttonpress and closes the program
    quitbutton = Button(homeframe,text='Quit',command=quitwindow).grid(row=0,column=3)
    #loop root over these defined objects to keep it displayed until something is done to it
    root.mainloop()

if(__name__=='__main__'):
    root = Tk() #define root as a tkinter window
    root.title('Manchester Harvard Reference Generator') #give the window this title string
    root.state('zoomed')
    elementscount=2
    main() #cal the main function
        
